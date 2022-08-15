# 执行机构：At home at college
# 出 品 人：🌫⭐※
# 开发时间:2021/10/12  20:16
'''1获取目标
2翻译
3生成新文件  '''
from docx import Document
import  os
import requests

url = "http://fanyi.youdao.com/translate"       #有道接口
#加载路径
path2='F:\\wwx使用\\buckup\\PycharmProjects\\实用工具\\翻译\\测试\\源文件'#指定输入目录
path3='F:\\wwx使用\\buckup\\PycharmProjects\\实用工具\\翻译\\测试\\生成目标'#输出目录
#path=os.path.abspath(os.curdir) #加载当前目录
#path1=path+'\\'+'源文件'
#os.chdir(path2) # 切换 工作位置
file=Document(path2+'/'+"1111.docx")  #加载文档(本代码段)注意文档格式只能读取docx
L=str(len(file.paragraphs))

#读取文件
a=0
L13=[]
while a < int(L):     #Len 加载文档长度（行数）
   s3=[file.paragraphs[int(a)].style,file.paragraphs[int(a)].text] # 段落样式(第几段)   返回指定行内容
   if s3[1]=='':
       L13.append(s3[1])
       aragraph = file.add_paragraph(L13[int(a)], style=s3[0])
       a += 1
       file.save(path3 + '/' + 'test.docx')  # 保存
       print('不翻译后保存')
   else:
       data = {'doctype': 'json', 'type': 'AUTO', 'i': s3[1]}
       r = requests.get(url, params=data)
       result = r.json()
       ss = result['translateResult'][0][0]['tgt']
       L13.append(ss)
       aragraph = file.add_paragraph(L13[int(a)], style=s3[0])
       a += 1
       file.save(path3 + '/' + 'test.docx')  # 保存
       print('翻译后保存')
#中英文自动翻译



