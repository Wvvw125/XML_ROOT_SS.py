# 执行机构：At home at college
# 出 品 人：🌫⭐※
# 开发时间:2021/10/12  20:16
'''1获取目标
2翻译
3生成新文件'''
from win32com import client as wc
from docx import Document
from docx.oxml.ns import qn
import  os
import docx
import  re

word = wc.Dispatch("Word.Application")
doc = word.Documents.Open('F:/wwx使用/buckup/PycharmProjects/实用工具/翻译/测试/源文件/1111.doc')
doc.SaveAs('F:/wwx使用/buckup/PycharmProjects/实用工具/翻译/测试/源文件/1111.doc', 2)  #12 为docx
doc.Close()
word.Quit()
# path=os.path.abspath(os.curdir) 加载当前目录
# print (path)
#file=Document("F:/wwx使用/buckup/PycharmProjects/实用工具/翻译/测试/源文件")
# path='F:\wwx使用\buckup\PycharmProjects\实用工具\翻译\测试\源文件'
# os.chdir(path) 切换 工作位置
# print('0'+str(len(file.par)))
#s=document.paragraphs[0].text
#print(s)
for p in doc.paragraphs:   # 读取标题
 if p.style.name=='Heading 1':
     print(p.text)
for p in doc.paragraphs:
  if re.match("^Heading \d+$",p.style.name):
     print(p.text)