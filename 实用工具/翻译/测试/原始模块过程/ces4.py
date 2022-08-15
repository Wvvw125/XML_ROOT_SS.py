# 执行机构：At home at college
# 出 品 人：🌫⭐※
# 开发时间:2021/10/12  20:16
'''1获取目标
2翻译
3生成新文件'''
from win32com import client as wc
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
import  os
import docx
import translate
import re

#锁定加载路径
path=os.path.abspath(os.curdir) #加载当前目录
print (path)
print('----')
path1=path+'\\'+'源文件'
print(path1)
os.chdir(path1) # 切换 工作位置

file=Document("1111.docx")
# document = Document()  #创建基于默认“模板”的空白文档
# document = Document('d:/test.docx')  #打开文档

paragraph =file.add_paragraph('段落3')  #在尾部添加段落
# paragraph.style = 'Heding1'
# paragraph =file.add_paragraph('段落4',style = 'Heding1')  #添加段落--带段落样式

print(file)
L=str(len(file.paragraphs)) #加载文档长度
print(L)
s1=file.paragraphs
#print('段落数为：'+str(len(file.paragraphs)))
#for p0 in s1:   #输出每一段内容
#    print(p0.text)
a=0
L1=[]
L2=[]
while a < int(L):
   s3 =s1[int(a)].style  # 段落样式(第几段)
   s = file.paragraphs[int(a)].text  # 返回指定行内容
   print(s)
   print(s3)
   a+=1
   L1.append(s)
   L2.append(s3)
print(L1)
print(L2)
file.save(path+'\\'+'生成目标'+'\\'+'test.docx')  #保存

