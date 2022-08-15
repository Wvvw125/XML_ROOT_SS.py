# 执行机构：At home at college
# 出 品 人：🌫⭐※
# 开发时间:2021/10/12  20:16
'''1获取目标
2翻译
3生成新文件'''
from win32com import client as wc
from docx import Document
from docx.oxml.ns import qn
import  os
import docx
import translate

#锁定加载路径
path=os.path.abspath(os.curdir) #加载当前目录
print (path)
print('----')
path1=path+'\\'+'源文件'
print(path1)
os.chdir(path1) # 切换 工作位置

file=Document("1111.docx")
print(file)
# path='F:\wwx使用\buckup\PycharmProjects\实用工具\翻译\测试\源文件'
print('段落数为：'+str(len(file.paragraphs)))

for para in file.paragraphs:   #输出每一段内容
    print(para.text)

for i in range(len(file.paragraphs)):  # 段落编号内容
    print('第'+str(i)+'段内容是：'+file.paragraphs[i].text)