# 执行机构：At home at college
# 出 品 人：🌫⭐※
# 开发时间:2021/10/12  20:16
'''1获取目标
2翻译
3生成新文件'''
from docx import Document
import  os
from translate import Translator

#锁定加载路径
path2='F:\\wwx使用\\buckup\\PycharmProjects\\实用工具\\翻译\\测试\\源文件'#指定输入目录
path3='F:\\wwx使用\\buckup\\PycharmProjects\\实用工具\\翻译\\测试\\生成目标'#输出目录
#path=os.path.abspath(os.curdir) #加载当前目录
#path1=path+'\\'+'源文件'
os.chdir(path2) # 切换 工作位置
file=Document("1111.docx")    #(本代码段)注意文档格式只能读取docx
L=str(len(file.paragraphs)) #加载文档长度（行数）
#读取文件
a=0
L13=[]
while a < int(L):
   s3 =file.paragraphs[int(a)].style  # 段落样式(第几段)
   s = file.paragraphs[int(a)].text  # 返回指定行内容
   translator = Translator(from_lang="chinese", to_lang="english")
   L13.append(translator.translate(s))
   aragraph = file.add_paragraph(L13[int(a)], style=s3)
   a+=1
#写入文件
file.save(path3+"/"+'text.docx')  #保存
print('翻译后保存完成')

